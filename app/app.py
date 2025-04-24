import streamlit as st
import PyPDF2
import docx
import os
import io
import base64
import sys
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import anthropic
from dotenv import load_dotenv
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from data.rag_summarizer import enhance_summary_with_rag
from data.utils import log_feedback

# Session state initialisatie 
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'context_info' not in st.session_state:
    st.session_state.context_info = None
if 'feedback' not in st.session_state:
    st.session_state.feedback = None
if 'has_generated' not in st.session_state:
    st.session_state.has_generated = False
if 'gebruikt_kennisbank_items' not in st.session_state:
    st.session_state.gebruikt_kennisbank_items = None    
if 'feedback_detail_step' not in st.session_state:
    st.session_state.feedback_detail_step = False
if 'feedback_complete' not in st.session_state:
    st.session_state.feedback_complete = False
# NIEUW: Chat-gerelateerde session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'original_text' not in st.session_state:
    st.session_state.original_text = None
if 'chat_view_active' not in st.session_state:
    st.session_state.chat_view_active = False

# Zorg dat Python ook de rootmap kan vinden
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from ai_summarizer import config
from ai_summarizer import doelgroepen_config
from ai_summarizer import avatars_config

# Laad API sleutel
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '..', '.env'))
anthropic_api_key = st.secrets["ANTHROPIC_API_KEY"]

# Configuratie
st.set_page_config(page_title="Document Summarizer", layout="wide")

# Styling
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stButton button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
    }
    .stTitle {
        font-size: 2.5rem;
        font-weight: bold;
        color: #2E2E2E;
    }
    .avatar-img {
        border-radius: 50%;
        width: 80px;
        height: 80px;
        object-fit: cover;
    }
    /* NIEUW: Chat styling */
    .chat-container {
        margin: 1rem 0;
        border-radius: 10px;
        padding: 1rem;
        background-color: #f8f9fa;
    }
    .chat-message-user {
        background-color: #e3f2fd;
        border-radius: 10px;
        padding: 0.8rem;
        margin: 0.5rem 0;
        text-align: right;
        margin-left: 20%;
    }
    .chat-message-assistant {
        background-color: #f0f0f0;
        border-radius: 10px;
        padding: 0.8rem;
        margin: 0.5rem 0;
        margin-right: 20%;
    }
</style>
""", unsafe_allow_html=True)

# Titel
st.title("Document Summarizer")
st.write("Upload je documenten, kies je voorkeuren, en ontvang een gepersonaliseerde samenvatting.")

# Aangepaste styling voor avatars met geforceerde uniforme afmetingen
st.markdown("""
<style>
    .avatar-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        text-align: center;
        margin-bottom: 10px;
    }
    .avatar-name {
        margin-top: 5px;
        font-weight: bold;
    }
    /* Deze CSS zorgt ervoor dat alle avatar-afbeeldingen dezelfde grootte hebben */
    .uniform-avatar {
        width: 80px !important;
        height: 80px !important;
        object-fit: cover !important;
        border-radius: 50% !important;
    }
</style>
""", unsafe_allow_html=True)

# Avatar selectie
st.markdown("### Kies je samenvattingsstijl")
cols = st.columns(len(avatars_config.AVATARS))

if "selected_avatar" not in st.session_state:
    st.session_state.selected_avatar = list(avatars_config.AVATARS.keys())[0]

for i, (key, avatar) in enumerate(avatars_config.AVATARS.items()):
    with cols[i]:
        avatar_img_path = os.path.join(avatars_config.AVATAR_IMAGES_PATH, f"{key}.jpg")
        try:
            # Laad de afbeelding en gebruik HTML om de grootte af te dwingen
            from PIL import Image
            import base64
            from io import BytesIO

            img = Image.open(avatar_img_path)
            buffered = BytesIO()
            img.save(buffered, format=img.format or "PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode()

            st.markdown(
                f"""
                <div class="avatar-container">
                    <img src="data:image/{img.format.lower() if img.format else 'png'};base64,{img_str}" class="uniform-avatar">
                    <div class="avatar-name">{avatar['name']}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        except Exception as e:
            st.error(f"Fout bij laden afbeelding: {str(e)}")
        
        # Als geen afbeelding gevonden is, toon placeholder
        if not avatar_found:
            st.markdown(f"""
            <div class="avatar-container">
                <div style="width:80px;height:80px;border-radius:50%;background-color:#f0f0f0;display:flex;justify-content:center;align-items:center;">📷</div>
                <div class="avatar-name">{avatar['name']}</div>
            </div>
            """, unsafe_allow_html=True)
        
        if st.button(f"Selecteer {avatar['name']}", key=f"avatar_{key}"):
            st.session_state.selected_avatar = key

selected_avatar = st.session_state.selected_avatar
avatar_info = avatars_config.AVATARS[selected_avatar]
st.success(f"🎯 Actieve avatar: {avatar_info['name']} - {avatar_info['description']}")

# Status van API-sleutel tonen
if not anthropic_api_key:
    st.error("⚠️ Geen API-sleutel gevonden. Zet je .env bestand in de rootmap.")
    api_key_input = st.text_input("Voer je Anthropic API-sleutel in:", type="password")
    if api_key_input:
        anthropic_api_key = api_key_input
        st.success("API-sleutel ingesteld!")

# Sidebar: voorkeuren en documentupload
with st.sidebar:
    st.header("Documentinvoer")
    uploaded_files = st.file_uploader("Upload PDF of Word bestanden", 
                                     type=["pdf", "docx", "doc"], 
                                     accept_multiple_files=True)

    st.header("Doelgroep")
    doelgroep = st.selectbox("Kies een doelgroep", list(doelgroepen_config.DOELGROEPEN.keys()))
    defaults = doelgroepen_config.DOELGROEPEN[doelgroep]

    st.header("Voorkeuren")
    summary_length = st.selectbox("Lengte samenvatting", config.SUMMARY_LENGTHS, index=config.SUMMARY_LENGTHS.index(defaults["summary_length"]))
    style = st.selectbox("Stijl", config.STYLES, index=config.STYLES.index(defaults["style"]))
    formality = st.selectbox("Formaliteit", config.FORMALITY_LEVELS, index=config.FORMALITY_LEVELS.index(defaults["formality"]))
    purpose = st.selectbox("Doel", config.PURPOSES)
    industry = st.selectbox("Branche", config.INDUSTRIES)
    horizon = st.selectbox("Beleggingshorizon", config.INVESTMENT_HORIZONS)
    risk_profile = st.selectbox("Risicoprofiel", config.RISK_PROFILES)
    output_format = st.radio("Output formaat", ["PDF", "Word"])

    st.header("Geavanceerde opties")
    use_rag = st.toggle("Gebruik kennisbank (RAG)", value=True, 
                        help="Schakel deze uit om samenvattingen zonder kennisbank te genereren")

# Functies voor document verwerking

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def process_documents(uploaded_files):
    combined_text = ""
    for uploaded_file in uploaded_files:
        bytes_data = uploaded_file.getvalue()
        temp_file = io.BytesIO(bytes_data)
        temp_file.name = uploaded_file.name
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension == ".pdf":
            text = extract_text_from_pdf(temp_file)
        elif file_extension in [".docx", ".doc"]:
            text = extract_text_from_docx(temp_file)
        else:
            st.warning(f"Bestandstype {file_extension} wordt niet ondersteund.")
            continue
        combined_text += f"\n\n--- Document: {uploaded_file.name} ---\n\n{text}"
    return combined_text

def create_pdf(text, filename="samenvatting.pdf"):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica", 12)
    y_position = height - 50
    lines = text.split('\n')
    for line in lines:
        if y_position < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y_position = height - 50
        words = line.split()
        current_line = ""
        for word in words:
            if c.stringWidth(current_line + word, "Helvetica", 12) < width - 100:
                current_line += word + " "
            else:
                c.drawString(50, y_position, current_line)
                y_position -= 15
                current_line = word + " "
        if current_line:
            c.drawString(50, y_position, current_line)
            y_position -= 15
    c.save()
    buffer.seek(0)
    return buffer

def create_docx(text, filename="samenvatting.docx"):
    doc = Document()
    doc.add_heading("Samenvatting", 0)
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def log_summary_result(original_text, summary, used_rag, metadata=None):
    """Log summary results for evaluation"""
    log_file = os.path.join(os.path.dirname(__file__), '..', 'data', 'summary_logs.csv')
    
    # Create file with headers if it doesn't exist
    if not os.path.exists(log_file):
        with open(log_file, 'w', newline='', encoding='utf-8') as f:
            f.write("timestamp,document_length,summary_length,used_rag,avatar,doelgroep\n")
    
    # Add log entry
    with open(log_file, 'a', newline='', encoding='utf-8') as f:
        metadata = metadata or {}
        log_entry = f"{datetime.now().isoformat()},{len(original_text)},{len(summary)},{used_rag},"
        log_entry += f"{metadata.get('avatar', 'unknown')},{metadata.get('doelgroep', 'unknown')}\n"
        f.write(log_entry)

def get_download_link(buffer, filename, format_type):
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    mime_type = "application/pdf" if format_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {filename}</a>'
    return href

def summarize_text(text, summary_length, style, formality, purpose, industry, horizon, risk_profile, doelgroep, use_rag=True):
    try:
        client = anthropic.Anthropic(api_key=anthropic_api_key)
        
        # Bepaal of we RAG gebruiken op basis van de toggle
        if use_rag:
            # Gebruik de RAG functie om de input te verrijken met relevante context
            enhanced_text, context_info = enhance_summary_with_rag(client, text)
        else:
            # Zonder RAG, gebruik de originele tekst
            enhanced_text = text
            context_info = "Kennisbank niet gebruikt"

        # Haal de system prompt op voor de geselecteerde avatar
        system = avatars_config.AVATARS[selected_avatar]["system_role"]

        # Verbeterde prompt voor RAG
        if use_rag and isinstance(context_info, dict) and 'relevante_items' in context_info:
            relevante_termen = context_info['relevante_items']
            if relevante_termen:
                eerste_prompt = f"""
                Vat onderstaand document samen. Structureer in categorieën zoals hoofdboodschap, feiten, cijfers, risico's en kansen.
                
                In dit document komen mogelijk de volgende financiële termen voor: {', '.join(relevante_termen)}
                Leg deze termen duidelijk uit wanneer ze in het document voorkomen.
                Markeer de uitleg van deze termen in je samenvatting tussen haakjes zodat ze duidelijk herkenbaar zijn.
                Bijvoorbeeld: "Het bedrijf rapporteerde een hogere EBITDA (Earnings Before Interest, Taxes, Depreciation and Amortization, een maatstaf voor operationele winstgevendheid)."

                [BEGIN DOCUMENT]
                {enhanced_text}
                [END DOCUMENT]
                """
            else:
                # Standaard prompt als er geen relevante termen zijn
                eerste_prompt = f"""
                Vat onderstaand document samen. Structureer in categorieën zoals hoofdboodschap, feiten, cijfers, risico's en kansen.

                [BEGIN DOCUMENT]
                {enhanced_text}
                [END DOCUMENT]
                """
        else:
            # Standaard prompt zonder RAG
            eerste_prompt = f"""
            Vat onderstaand document samen. Structureer in categorieën zoals hoofdboodschap, feiten, cijfers, risico's en kansen.

            [BEGIN DOCUMENT]
            {enhanced_text}
            [END DOCUMENT]
            """

        # Eerste call naar Claude
        initial_message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=3000,
            temperature=0.3,
            system="Je bent een efficiënte samenvattings-assistent. Je vat documenten samen in overzichtelijke secties.",
            messages=[{"role": "user", "content": eerste_prompt}]
        )
        initial_summary = initial_message.content[0].text

        # Tweede prompt: herschrijf in avatarstijl - behouden van je bestaande code
        
        if selected_avatar in ["annick", "erin"]:
            structuur_prompt = """
            Je krijgt een document. Vat dit samen op een vloeiende manier, zonder opsommingen of lijstjes. Gebruik jouw unieke stijl.
            Geef uitleg alsof je iemand persoonlijk iets vertelt. Geen bullets, geen stappenplan.

            Als je moeilijke termen tegenkomt, leg ze dan uit op een natuurlijke manier — bijvoorbeeld met een vergelijking of metafoor.
            Gebruik liever geen haakjes of encyclopedische definities, maar maak de uitleg onderdeel van je verhaal.
            """
        else:
            structuur_prompt = """  Je krijgt een document. Vat dit samen op een vloeiende manier. Gebruik jouw unieke stijl.
            Geef uitleg alsof je iemand persoonlijk iets vertelt.Gebruik cijfers indien relevant. Ook moeilijk jargon is zeker toegelaten
            
            
            Volg zorgvuldig deze stappen:
            1. Vat de hoofdboodschap kort samen.
            2. Benoem de belangrijkste feiten, cijfers en relevante marktinformatie.
            3. Vermeld expliciet eventuele risico's en opportuniteiten.
            4. Formuleer duidelijk, gestructureerd en aangepast aan de volgende voorkeuren:
                - Lengte: {summary_length}
                - Stijl: {style}
                - Formaliteit: {formality}. {config.FORMALITY_DESCRIPTIONS.get(formality, '')}
                - Doel: {purpose}. {config.PURPOSE_DESCRIPTIONS.get(purpose, '')}
                - Branche: {industry}. {config.INDUSTRY_DESCRIPTIONS.get(industry, '')}
                - Beleggingshorizon: {horizon}. {config.INVESTMENT_HORIZON_DESCRIPTIONS.get(horizon, '')}
                - Risicoprofiel: {risk_profile}. {config.RISK_PROFILE_DESCRIPTIONS.get(risk_profile, '')}
                - Doelgroep: {doelgroep}. {doelgroepen_config.DOELGROEP_DESCRIPTIONS.get(doelgroep, '')}
            """

        # Toevoegen van expliciete instructies voor financiële termen
        if use_rag and isinstance(context_info, dict) and 'relevante_items' in context_info:
            relevante_termen = context_info['relevante_items']
            if relevante_termen:
                structuur_prompt += f"""
                
                BELANGRIJK: In dit document komen de volgende financiële termen voor: {', '.join(relevante_termen)}
                Leg deze termen duidelijk uit wanneer ze in het document voorkomen.
                Markeer je uitleg tussen haakjes voor duidelijkheid, bijvoorbeeld: 
                "De EBITDA (Earnings Before Interest, Taxes, Depreciation and Amortization, een maatstaf voor operationele winstgevendheid) is gestegen."
                """
        
        prompt = f"""
        {structuur_prompt}
        
        [BEGIN DOCUMENT]
        {initial_summary}
        [END DOCUMENT]

        Maak een beknopte, coherente en begrijpelijke samenvatting zonder dubbelingen of overbodige informatie. Zorg ervoor dat de samenvatting relevant is voor de gekozen avatar, doelgroep en voorkeuren.
        """

        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4000,
            temperature=0.7,
            system=system,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return message.content[0].text, context_info
    except Exception as e:
        st.error(f"❌ Error generating summary: {e}")
        return "Er is een fout opgetreden bij het genereren van de samenvatting. Probeer het opnieuw.","Error"

# NIEUW: Verbeterde Chat functionaliteit met avatar stijl
def generate_chat_response(query, document_text, summary, chat_history, avatar_system, use_rag=True):
    try:
        client = anthropic.Anthropic(api_key=anthropic_api_key)
        
        # Bouw chat history formaat voor Claude API
        formatted_history = []
        for entry in chat_history:
            formatted_history.append({"role": entry["role"], "content": entry["content"]})
        
        # Controleer of we RAG moeten gebruiken voor meer context
        context_chunks = document_text
        if use_rag:
            try:
                # Gebruik de standaard enhance_summary_with_rag zonder extra parameters
                enhanced_text, context_info = enhance_summary_with_rag(client, document_text)
                
                # Controleer of we relevante inhoud hebben
                if isinstance(context_info, dict) and 'relevante_context' in context_info:
                    context_chunks = context_info['relevante_context']
                elif isinstance(context_info, dict) and 'relevante_items' in context_info:
                    # Gebruik de relevante items als er geen directe context is
                    context_chunks = document_text + "\n\nRelevante context: " + ", ".join(context_info['relevante_items'])
            except Exception as e:
                st.info(f"Info: RAG werkt beperkt voor chat vragen. Valt terug op volledige document.")
                # Valt terug op het volledige document als RAG faalt
                pass
        
        # Bouw system prompt met behoud van avatar persoonlijkheid
        system_prompt = f"""
        {avatar_system}
        
        Je bent een persoonlijke financiële assistent met een unieke communicatiestijl.
        
        BELANGRIJKE INSTRUCTIES:
        1. Behoud ALTIJD je persoonlijke communicatiestijl wanneer je antwoord geeft.
        2. Baseer je antwoord UITSLUITEND op informatie uit het document.
        3. Als de vraag niet beantwoord kan worden met de informatie in het document, zeg dit op een vriendelijke manier in je eigen stijl.
        4. Gebruik geen externe kennis die niet in het document staat.
        5. Je mag het document nooit rechtstreeks citeren zonder het aan te passen naar je eigen stijl.
        
        DISCLAIMER: Benadruk waar nodig dat je antwoorden gebaseerd zijn op het document en geen persoonlijk financieel advies zijn.
        """
        
        # Voorbereiding van de context voor het model
        context = f"""
        DOCUMENT SAMENVATTING:
        {summary}
        
        RELEVANTE DELEN UIT HET DOCUMENT:
        {context_chunks}
        
        SPECIFIEKE VRAAG VAN DE GEBRUIKER OVER HET DOCUMENT:
        {query}
        """
        
        # Eerste stap: Antwoord genereren met avatar-persoonlijkheid
        initial_response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=1500,
            temperature=0.5,  # Iets hogere temperatuur voor meer persoonlijkheid
            system=system_prompt,
            messages=[
                {"role": "user", "content": f"Hier is wat context voor mijn vraag: {context}"},
                *formatted_history,
                {"role": "user", "content": query}
            ]
        )
        response_text = initial_response.content[0].text
        
        # Tweede stap: Validatie / fact-checking
        validation_prompt = f"""
        Je bent een strikte validator die controleert of antwoorden over financiële documenten accuraat zijn.
        
        Evalueer het volgende antwoord:
        
        VRAAG: {query}
        
        GEGEVEN ANTWOORD: {response_text}
        
        DOCUMENT CONTEXT: {context}
        
        Controleer alleen op feitelijke onjuistheden of beweringen die niet in het document staan.
        Beoordeel NIET op stijl, toon of persoonlijkheid - die aspecten moeten behouden blijven.
        
        Als er feitelijke problemen zijn, geef aan welke claims niet kloppen. Als alle feiten correct zijn, antwoord alleen met "GOEDGEKEURD".
        """
        
        validation_check = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=500,
            temperature=0.1,
            messages=[{"role": "user", "content": validation_prompt}]
        )
        
        validation_result = validation_check.content[0].text
        
        # Als het antwoord niet is goedgekeurd, filter het maar behoud de stijl
        if "GOEDGEKEURD" not in validation_result.upper():
            # Herformuleer om problematische delen te verwijderen, maar behoud stijl
            correction_prompt = f"""
            Je antwoord bevat mogelijk feitelijke informatie die niet in het document staat.
            
            Probleem: {validation_result}
            
            Herformuleer je antwoord met dezelfde persoonlijke stijl en toon, maar:
            1. Verwijder of corrigeer claims die niet door het document worden ondersteund
            2. Behoud je oorspronkelijke stem, stijl en persoonlijkheid
            3. Wees explicieter wanneer je aangeeft dat iets niet in het document staat
            
            Oorspronkelijke vraag: {query}
            Oorspronkelijk antwoord in jouw stijl: {response_text}
            """
            
            correction = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1500,
                temperature=0.4,  # Behoud enige creativiteit voor de stijl
                system=system_prompt,  # Hergebruik dezelfde avatar-instructies
                messages=[{"role": "user", "content": correction_prompt}]
            )
            
            final_response = correction.content[0].text
        else:
            final_response = response_text
            
        return final_response
    except Exception as e:
        st.error(f"Error genereren chatantwoord: {str(e)}")
        return "Er is een fout opgetreden bij het genereren van een antwoord. Probeer het opnieuw."
# Functie voor chat geschiedenis loggen
def log_chat_interaction(question, answer, document_info=None):
    """Log chat interactions for evaluation"""
    log_file = os.path.join(os.path.dirname(__file__), '..', 'data', 'chat_logs.csv')
    
    # Create file with headers if it doesn't exist
    if not os.path.exists(log_file):
        with open(log_file, 'w', newline='', encoding='utf-8') as f:
            f.write("timestamp,question_length,answer_length,avatar,rag_used\n")
    
    # Add log entry
    with open(log_file, 'a', newline='', encoding='utf-8') as f:
        document_info = document_info or {}
        log_entry = f"{datetime.now().isoformat()},{len(question)},{len(answer)},"
        log_entry += f"{document_info.get('avatar', 'unknown')},{document_info.get('rag_used', True)}\n"
        f.write(log_entry)

# Hoofdsectie
if uploaded_files:
    st.header("Verwerking")
    with st.spinner("Documenten verwerken..."):
        combined_text = process_documents(uploaded_files)
        # NIEUW: Bewaar de originele tekst voor de chatbot
        st.session_state.original_text = combined_text
    st.success(f"Succesvol {len(uploaded_files)} document(en) verwerkt.")
    with st.expander("Bekijk ruwe tekst"):
        st.text_area("Gecombineerde tekst", combined_text, height=200)
    
    if st.button("Genereer samenvatting"):
        if not anthropic_api_key:
            st.error("⚠️ API-sleutel ontbreekt. Voer deze in om samenvattingen te genereren.")
        else:
            with st.spinner("Samenvatting genereren..."):
                summary, context_info = summarize_text(
                    combined_text,
                    summary_length,
                    style,
                    formality,
                    purpose,
                    industry,
                    horizon,         
                    risk_profile,    
                    doelgroep,
                    use_rag            # Geef de toggle waarde door
                )
            
            # Bewaar in session state
            st.session_state.summary = summary
            st.session_state.context_info = context_info
            st.session_state.has_generated = True
            
            # Reset chat historiek als nieuwe samenvatting is gemaakt
            st.session_state.chat_history = []

            # Reset feedback state voor nieuwe samenvatting
            st.session_state.feedback = None
            st.session_state.feedback_detail_step = False
            st.session_state.feedback_complete = False

            log_summary_result(
                combined_text, 
                summary, 
                use_rag, 
                metadata={
                    'avatar': selected_avatar,
                    'doelgroep': doelgroep
                }
            )
                
# Na de knop, controleer of er een samenvatting is om te tonen
if st.session_state.summary is not None:
    st.header("Samenvatting")
    
    # NIEUW: Tabbladen voor samenvatting en chat
    tab1, tab2 = st.tabs(["Samenvatting", f"Stel {avatar_info['name']} bijkomende vragen over het ingeladen document"])
   
    with tab1:
        st.write(st.session_state.summary)
        
        # Optioneel: Toon informatie over het RAG-gebruik
        with st.expander("Technische details"):
            if use_rag:
                st.info("Deze samenvatting is gegenereerd met behulp van de kennisbank (RAG).")
                
                # Controleer of context_info een dictionary is met relevante items
                if isinstance(st.session_state.context_info, dict) and 'relevante_items' in st.session_state.context_info:
                    items = st.session_state.context_info['relevante_items']
                    if items:
                        st.subheader("Gebruikte kennisbank items:")
                        for item in items:
                            st.markdown(f"- **{item}**")
                        
                        # Toon de algemene toelichting
                        if 'toelichting' in st.session_state.context_info:
                            st.write(st.session_state.context_info['toelichting'])
                    else:
                        st.write("Er zijn geen relevante kennisbank-items gevonden voor dit document.")
                else:
                    st.write("Gebruikte context:", st.session_state.context_info)
            else:
                st.info("Deze samenvatting is gegenereerd zonder gebruik van de kennisbank.")
                
        # Hier de feedback knoppen toevoegen
        if st.session_state.feedback is None:
            st.write("Was deze samenvatting nuttig?")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("👍 Ja", key="btn_yes"):
                    st.session_state.feedback = "positief"
                    st.session_state.feedback_detail_step = True
                    # Log initiële feedback
                    log_feedback(
                        feedback_type="positief",
                        feedback_details="initial_response",
                        document_info={
                            'length': len(combined_text),
                            'avatar': selected_avatar,
                            'rag_used': use_rag
                        }
                    )
                    st.rerun()
            with col2:
                if st.button("👎 Nee", key="btn_no"):
                    st.session_state.feedback = "negatief"
                    st.session_state.feedback_detail_step = True
                    # Log initiële feedback
                    log_feedback(
                        feedback_type="negatief",
                        feedback_details="initial_response",
                        document_info={
                            'length': len(combined_text),
                            'avatar': selected_avatar,
                            'rag_used': use_rag
                        }
                    )
                    st.rerun()
        # Toon het gedetailleerde feedbackformulier ALLEEN als we in die stap zijn én nog niet klaar zijn
        elif st.session_state.get('feedback_detail_step', False) and not st.session_state.get('feedback_complete', False):
            # Toon een formulier voor meer details
            if st.session_state.feedback == "positief":
                st.success("Fijn dat je de samenvatting nuttig vond!")
                feedback_detail = st.radio(
                    "Wat vond je het meest waardevol?",
                    ["De duidelijkheid", "De volledigheid", "De uitleg van termen", "De lengte was perfect", "Anders"]
                )
            else:
                st.warning("Bedankt voor je feedback. Help ons verbeteren!")
                feedback_detail = st.radio(
                    "Wat kon beter?",
                    ["Te lang", "Te kort", "Miste belangrijke informatie", "Verwarrende uitleg", "Kennisbank werd niet goed gebruikt", "Anders"]
                )
            
            other_feedback = ""
            if feedback_detail == "Anders":
                other_feedback = st.text_area("Vertel ons meer:")
            
            if st.button("Verstuur feedback"):
                # Log gedetailleerde feedback
                detail_to_log = other_feedback if feedback_detail == "Anders" else feedback_detail
                log_feedback(
                    feedback_type=st.session_state.feedback,
                    feedback_details=detail_to_log,
                    document_info={
                        'length': len(combined_text),
                        'avatar': selected_avatar,
                        'rag_used': use_rag
                    }
                )
                st.session_state.feedback_complete = True
                st.rerun()
        # Het formulier is ingevuld en de feedback is compleet - toon alleen het bedankbericht
        elif st.session_state.get('feedback_complete', False):
            # Toon een bericht dat feedback is ontvangen
            if st.session_state.feedback == "positief":
                st.success("Bedankt voor je feedback! We zijn blij dat de samenvatting nuttig was.")
            else:
                st.info("Bedankt voor je feedback. We gebruiken deze om onze tool te verbeteren.")

            # Toon download opties
            st.header("Download opties")
            if output_format == "PDF":
                pdf_buffer = create_pdf(st.session_state.summary)
                st.markdown(get_download_link(pdf_buffer, "samenvatting.pdf", "pdf"), unsafe_allow_html=True)
            else:
                docx_buffer = create_docx(st.session_state.summary)
                st.markdown(get_download_link(docx_buffer, "samenvatting.docx", "docx"), unsafe_allow_html=True)
    
    # NIEUW: Chat interface in tweede tab
    with tab2:
        st.write("Stel vragen over het document en krijg antwoorden in de stijl van de gekozen avatar.")
        st.write("De chatbot geeft alleen antwoorden gebaseerd op de inhoud van het document.")
        
        # Waarschuwing/disclaimer
        st.warning("⚠️ Let op: Deze chatbot geeft alleen informatie uit het document. Het geeft geen persoonlijk financieel advies en vervangt niet het advies van een financieel professional.")
        
        # Toon chat geschiedenis
        st.subheader("Gesprek")
        chat_container = st.container()
        
        with chat_container:
            for message in st.session_state.chat_history:
                if message["role"] == "user":
                    st.markdown(f"<div class='chat-message-user'>{message['content']}</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='chat-message-assistant'>{message['content']}</div>", unsafe_allow_html=True)
        
        # Chat input
        user_question = st.text_input("Stel een vraag over het document:", key="chat_input")
        
        if st.button("Verstuur vraag", key="send_question"):
            if user_question:
                # Toon de vraag
                st.session_state.chat_history.append({"role": "user", "content": user_question})
                
                with st.spinner("Antwoord genereren..."):
                    # Haal system prompt op van avatar
                    avatar_system = avatars_config.AVATARS[selected_avatar]["system_role"]
                    
                    # Genereer antwoord
                    response = generate_chat_response(
                        user_question,
                        st.session_state.original_text,
                        st.session_state.summary,
                        st.session_state.chat_history,
                        avatar_system,
                        use_rag
                    )
                    
                    # Voeg toe aan chat history
                    st.session_state.chat_history.append({"role": "assistant", "content": response})
                    
                    # Log de interactie
                    log_chat_interaction(
                        user_question, 
                        response, 
                        document_info={
                            'avatar': selected_avatar,
                            'rag_used': use_rag
                        }
                    )
                
                # Clear the input
                st.rerun()
else:
    st.info("Upload één of meerdere documenten om te beginnen.")

# Extra informatie
st.sidebar.markdown("---")
st.sidebar.markdown("### Over deze app")
st.sidebar.write("Deze app gebruikt AI om documenten samen te vatten en te optimaliseren volgens jouw voorkeuren. Je kunt ook chatten met het document om specifieke vragen te stellen.")
